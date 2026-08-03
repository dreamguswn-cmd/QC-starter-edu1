from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\aws_project\QC-starter-edu1-live\public\downloads\Yoo_Hyunju_Resume.docx"
FONT = "맑은 고딕"
INK = RGBColor(11, 23, 40)
BLUE = RGBColor(18, 104, 232)
MUTED = RGBColor(82, 96, 117)


def font(run, size=10.5, bold=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run


def add_link(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1268E8")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def para(doc, text="", size=10.5, bold=False, color=INK, after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    font(p.add_run(text), size, bold, color)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 11.5, True, BLUE)
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "3")
    border.set(qn("w:color"), "D6E1EA")
    bottom.append(border)
    p._p.get_or_add_pPr().append(bottom)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    font(p.add_run(text), 9.6)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
styles["Normal"].font.name = FONT
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
styles["Normal"].font.size = Pt(10.5)

p = para(doc, "유현주", 27, True, INK, 1)
p.paragraph_format.space_before = Pt(0)
para(doc, "AI Software QA Engineer 지원", 13, True, BLUE, 5)
para(doc, "AI Software QA 실무 경력 없음 · 교육 및 프로젝트 기반 지원", 9.5, True, MUTED, 8)

contact = doc.add_paragraph()
contact.paragraph_format.space_after = Pt(10)
font(contact.add_run("EMAIL  "), 8.8, True, MUTED)
add_link(contact, "aa01057559209@gmail.com", "mailto:aa01057559209@gmail.com")
font(contact.add_run("    GITHUB  "), 8.8, True, MUTED)
add_link(contact, "github.com/dreamguswn-cmd", "https://github.com/dreamguswn-cmd")
font(contact.add_run("    PORTFOLIO  "), 8.8, True, MUTED)
add_link(contact, "dreamguswn-cmd.github.io/QC-starter-edu1/", "https://dreamguswn-cmd.github.io/QC-starter-edu1/")

heading(doc, "PROFILE")
para(doc, "임상병리사로 약 8년간 검사 결과의 정확성, 재현성, 이상치 확인과 장비 품질관리를 수행했습니다. 현재 대우능력개발원 AI 기반 소프트웨어 테스터(QA) 및 모니터링 실무 과정에서 기능·API 테스트, AI 응답 품질평가, 자동화, 성능 검증, 모니터링과 AWS 장애복구를 학습하고 프로젝트로 증빙하고 있습니다. 의료 QC 경험을 소프트웨어 품질 검증 역량으로 확장해 신입 AI Software QA 직무에 지원합니다.", 10.2, False, INK, 7)

heading(doc, "QA EDUCATION")
para(doc, "대우능력개발원  |  AI 기반 소프트웨어 테스터(QA) 및 모니터링 실무 과정", 10.5, True, INK, 2)
para(doc, "2026.05.27 - 2026.08.07  |  수강 중", 9.2, False, MUTED, 4)
bullet(doc, "기능·API·회귀 테스트, 테스트 케이스와 결함 재현 기록")
bullet(doc, "pytest 자동화, k6 성능 검증, Prometheus·Grafana 모니터링")
bullet(doc, "RAG·LLM Judge 기반 AI 응답 품질평가 및 답변 개선")
bullet(doc, "AWS EC2·S3·CloudTrail 환경 구축, 장애 재현·복구·보안 점검")

heading(doc, "SELECTED QA PROJECTS")
projects = [
    ("RAG 챗봇 자동 평가 및 답변 개선  |  개인", "평가 기준 설계, 응답 5건 자동 평가, 감점 원인 분석, Correction Agent 개선과 재검증을 전 과정 수행. 평균 4.20/5."),
    ("AWS 웹 서버 장애 재현 및 복구 검증  |  개인", "EC2·Apache 구축부터 장애 시나리오, 원인 확인, 서비스 복구, S3·CloudTrail 보안 재검증까지 수행. 복구 평가 100점."),
    ("VOC Improve 멀티 에이전트 QA  |  팀", "Agent 결과 테스트 실행, AI 품질 결과 검토, Release Gate 배포 판정 자료와 최종 증빙 정리를 담당."),
    ("AI Agent 품질관리·운영 모니터링  |  교육 프로젝트", "pytest 기능 테스트와 k6 부하 테스트를 실행하고 Prometheus 지표를 Grafana 대시보드에서 검증."),
]
for title, detail in projects:
    para(doc, title, 10.1, True, INK, 1)
    para(doc, detail, 9.4, False, MUTED, 5)

heading(doc, "SKILLS - 교육 프로젝트 수행 근거 기준")
skills = [
    ("Testing", "7/10", "테스트 기준 수립, 기능·회귀·API 검증과 재테스트 수행"),
    ("AI Quality", "7/10", "RAG·LLM Judge 평가 기준 설계와 답변 개선·재검증 수행"),
    ("Automation", "6/10", "pytest 기반 반복 검증과 결과 보고 적용"),
    ("Performance", "5/10", "k6 부하 조건과 응답 시간·오류율 검증 실습"),
    ("Monitoring", "6/10", "Prometheus 지표 수집과 Grafana 대시보드 확인"),
    ("Cloud & Ops", "6/10", "AWS 서버 구축·장애복구·S3 보안 점검 개인 과제"),
]
for title, score, note in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(f"{title}  {score}  "), 9.7, True, BLUE)
    font(p.add_run(note), 9.3, False, INK)

heading(doc, "TRANSFERABLE QUALITY EXPERIENCE - 이전 직무")
para(doc, "AI Software QA 회사 경력은 없으며, 의료기관에서 약 8년간 수행한 품질관리 경험 중 지원 직무와 연결되는 역량만 요약했습니다.", 9.3, False, MUTED, 6)
bullet(doc, "매일 정도관리(QC)를 실행하고 기준 이탈 시 원인을 확인한 경험")
bullet(doc, "검사 결과의 정확성·재현성을 확인하고 레퍼런스를 재설정·재검증한 경험")
bullet(doc, "검사 장비 품질관리(QA), 결과 보고, 협업과 사내 교육 경험")
bullet(doc, "의료 결과가 진단과 치료로 이어지는 환경에서 근거와 기록을 중시한 경험")

heading(doc, "EDUCATION & CERTIFICATIONS")
para(doc, "신흥대학교(현 신한대학교)  |  임상병리과 졸업 · 보건전문학사", 10, True, INK, 5)
para(doc, "임상병리사 면허 · MOS Word Expert · MOS Excel Expert · MOS PowerPoint · MOS Access", 9.7, False, INK, 4)

heading(doc, "TOOLS")
para(doc, "Python · pytest · k6 · Prometheus · Grafana · AWS EC2 · S3 · CloudTrail · Docker · GitHub Actions · Jupyter · Streamlit · MS Office", 9.7, False, INK, 7)

heading(doc, "PORTFOLIO EVIDENCE")
para(doc, "각 프로젝트의 문제 정의, 테스트 설계·실행, 결함·개선, 재검증 결과와 증빙은 포트폴리오와 GitHub 저장소에서 확인할 수 있습니다.", 9.7, False, INK, 4)
p = doc.add_paragraph()
add_link(p, "포트폴리오 프로젝트 자세히 보기", "https://dreamguswn-cmd.github.io/QC-starter-edu1/#projects")
font(p.add_run("    "), 9)
add_link(p, "GitHub QA 문서 보기", "https://github.com/dreamguswn-cmd/QC-starter-edu1/tree/main/docs")

for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("유현주 · AI Software QA Engineer 지원 이력서"), 8, False, MUTED)

doc.core_properties.title = "유현주 AI Software QA Engineer 지원 이력서"
doc.core_properties.subject = "교육 및 프로젝트 기반 신입 QA 지원"
doc.core_properties.author = "유현주"
doc.save(OUT)
print(OUT)
